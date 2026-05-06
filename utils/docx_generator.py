from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO


THEMES = {
    'classic':   {'primary': (0x1E, 0x3A, 0x5F), 'accent': (0x25, 0x63, 0xEB), 'muted': (0x64, 0x74, 0x8B)},
    'modern':    {'primary': (0x11, 0x18, 0x27), 'accent': (0x7C, 0x3A, 0xED), 'muted': (0x6B, 0x72, 0x80)},
    'minimal':   {'primary': (0x00, 0x00, 0x00), 'accent': (0x37, 0x41, 0x51), 'muted': (0x6B, 0x72, 0x80)},
    'executive': {'primary': (0x1A, 0x1A, 0x1A), 'accent': (0xB4, 0x53, 0x09), 'muted': (0x6B, 0x72, 0x80)},
}


def _rgb(tup):
    return RGBColor(*tup)


def _add_border_bottom(para, color_hex='E5E7EB', size='4'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


class DocxGenerator:
    def generate(self, data):
        doc = Document()

        # Remove default styles and set narrow margins
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(2.2)
            section.right_margin  = Cm(2.2)

        # Remove spacing from Normal style
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after  = Pt(2)

        theme = THEMES.get(data.get('template', 'classic'), THEMES['classic'])
        primary = _rgb(theme['primary'])
        accent  = _rgb(theme['accent'])
        muted   = _rgb(theme['muted'])
        personal = data.get('personal', {})

        # ── NAME ────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(personal.get('name', 'Nama Anda'))
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = primary

        # ── CONTACT ─────────────────────────────────────────────────
        parts = []
        for key in ['email', 'phone', 'location']:
            v = personal.get(key, '').strip()
            if v:
                parts.append(v)
        linkedin = personal.get('linkedin', '').strip()
        if linkedin:
            li = linkedin.replace('https://www.linkedin.com/in/', '').replace('https://linkedin.com/in/', '').strip('/')
            parts.append(f'linkedin.com/in/{li}')
        if personal.get('website', '').strip():
            parts.append(personal['website'].strip())

        if parts:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            run2 = p2.add_run(' | '.join(parts))
            run2.font.size = Pt(8.5)
            run2.font.color.rgb = muted

        # Heavy divider after header
        div_para = doc.add_paragraph()
        _add_border_bottom(div_para, color_hex=('%02X%02X%02X' % theme['accent']), size='8')
        div_para.paragraph_format.space_after = Pt(4)

        def section_header(title):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = accent
            _add_border_bottom(p)

        def add_title_row(title_text, date_text=''):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(1)
            r1 = p.add_run(title_text)
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = primary
            if date_text:
                r2 = p.add_run(f'  {date_text}')
                r2.font.size = Pt(8.5)
                r2.font.color.rgb = muted

        def add_sub(text):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = muted

        def add_body(text):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)

        def add_bullet(text):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent  = Cm(0.5)
            run = p.add_run(text.strip())
            run.font.size = Pt(9)

        # ── SUMMARY ─────────────────────────────────────────────────
        summary = data.get('summary', '').strip()
        if summary:
            section_header('Professional Summary')
            add_body(summary)

        # ── EXPERIENCE ──────────────────────────────────────────────
        experience = data.get('experience', [])
        if experience:
            section_header('Work Experience')
            for exp in experience:
                end = 'Present' if exp.get('current') else (exp.get('end_date') or 'Present')
                date_range = f"{exp.get('start_date', '')} – {end}".strip(' –')
                add_title_row(exp.get('title', ''), date_range)

                company_line = exp.get('company', '')
                if exp.get('location'):
                    company_line += f" · {exp['location']}"
                if company_line:
                    add_sub(company_line)

                bullets = exp.get('bullets', [])
                if not bullets and exp.get('description'):
                    bullets = [exp['description']]
                for b in bullets:
                    if b.strip():
                        add_bullet(b)

        # ── EDUCATION ───────────────────────────────────────────────
        education = data.get('education', [])
        if education:
            section_header('Education')
            for edu in education:
                degree = edu.get('degree', '')
                if edu.get('field'):
                    degree += f" in {edu['field']}"
                start = edu.get('start_date', '')
                end   = edu.get('end_date', '')
                date_range = f'{start} – {end}'.strip(' –') if (start or end) else ''
                add_title_row(degree, date_range)
                add_sub(edu.get('school', ''))
                if edu.get('gpa'):
                    add_body(f"GPA: {edu['gpa']}")
                if edu.get('activities'):
                    add_body(edu['activities'])

        # ── SKILLS ──────────────────────────────────────────────────
        skills = data.get('skills', [])
        skill_cats = data.get('skill_categories', {})
        if skills or skill_cats:
            section_header('Skills')
            if skill_cats:
                for cat, cat_skills in skill_cats.items():
                    if cat_skills:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(2)
                        r = p.add_run(f'{cat}: ')
                        r.bold = True
                        r.font.size = Pt(9)
                        r2 = p.add_run(' · '.join(cat_skills))
                        r2.font.size = Pt(9)
            else:
                add_body(' · '.join(skills))

        # ── CERTIFICATIONS ──────────────────────────────────────────
        certifications = data.get('certifications', [])
        if certifications:
            section_header('Certifications')
            for cert in certifications:
                line = cert.get('name', '')
                if cert.get('authority'):
                    line += f" – {cert['authority']}"
                if cert.get('date'):
                    line += f" ({cert['date']})"
                add_bullet(line)

        # ── PROJECTS ────────────────────────────────────────────────
        projects = data.get('projects', [])
        if projects:
            section_header('Projects')
            for proj in projects:
                add_title_row(proj.get('name', ''))
                if proj.get('technologies'):
                    add_sub(proj['technologies'])
                if proj.get('description'):
                    add_body(proj['description'])

        # ── LANGUAGES ───────────────────────────────────────────────
        languages = data.get('languages', [])
        if languages:
            section_header('Languages')
            lang_parts = []
            for lang in languages:
                part = lang.get('name', '')
                if lang.get('proficiency'):
                    part += f" ({lang['proficiency']})"
                if part:
                    lang_parts.append(part)
            add_body(' · '.join(lang_parts))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
